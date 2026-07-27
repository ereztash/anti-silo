from __future__ import annotations

import csv
import html
import json
import re
from dataclasses import dataclass
from pathlib import Path

from .extract_limits import (
    MAX_COLUMNS,
    MAX_EXTRACTED_CHARS,
    MAX_PDF_PAGES,
    MAX_ROWS,
    Budget,
    archive_refusal_note,
    archive_too_large,
    clamp,
    limit_note,
)
from .scanner import read_text


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    status: str = "complete"
    note: str = ""


def _extract_csv(path: Path) -> ExtractionResult:
    budget = Budget()
    truncated = False
    try:
        with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as f:
            reader = csv.reader(f)
            for idx, row in enumerate(reader):
                if idx >= MAX_ROWS:
                    truncated = True
                    break
                if len(row) > MAX_COLUMNS:
                    truncated = True
                cells = row[:MAX_COLUMNS]
                if not budget.add(" | ".join(cell.strip() for cell in cells)):
                    truncated = True
                    break
    except (csv.Error, OSError, ValueError) as exc:
        return ExtractionResult("", "failed", f"CSV could not be read: {exc}")
    return ExtractionResult(
        budget.text(),
        "truncated" if truncated else "complete",
        limit_note("CSV") if truncated else "",
    )


def _extract_json(path: Path) -> ExtractionResult:
    try:
        with path.open("r", encoding="utf-8-sig", errors="replace") as f:
            data = json.load(f)
    except (ValueError, OSError) as exc:
        return ExtractionResult("", "failed", f"JSON could not be parsed: {exc}")
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if len(text) > 20000:
        return ExtractionResult(text[:20000], "truncated", "JSON limited to the first 20,000 characters")
    return ExtractionResult(text)


def _strip_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _extract_docx(path: Path) -> ExtractionResult:
    try:
        import docx  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional package
        return ExtractionResult("", "failed", f"DOCX extraction unavailable: {exc}")
    oversized = archive_too_large(path)
    if oversized is not None:
        return ExtractionResult("", "failed", archive_refusal_note("DOCX", oversized))
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        return ExtractionResult("", "failed", f"DOCX could not be read (corrupt or not a real .docx): {exc}")
    budget = Budget()
    truncated = False

    def take(text: str) -> bool:
        """False once the budget is spent, so every caller stops the same way."""
        nonlocal truncated
        if not text.strip():
            return True
        if not budget.add(text):
            truncated = True
            return False
        return True

    # Paragraphs alone miss the shape most commercial documents actually use. A
    # rate card, fee schedule or term sheet is a TABLE, and `document.paragraphs`
    # never descends into one. Measured before this fix: two Word rate cards
    # whose every day rate differed (6400/7100, 4900/5450, 4200/4700, 3100/3450,
    # 2350/2600) each extracted to 39 characters - the heading - and were
    # reported `complete`. The near-duplicate detector then compared two nearly
    # empty strings, called them 100% identical, filed it `cleanup`, and advised
    # deleting one. Reading nothing while reporting success is worse than failing
    # to read: every other cap in this module truncates loudly, this path did not.
    for paragraph in document.paragraphs:
        if not take(paragraph.text):
            break
    for table in document.tables if not truncated else []:
        for row in table.rows:
            seen: set[int] = set()  # merged cells repeat one object across a row
            cells = [c for c in row.cells if id(c) not in seen and not seen.add(id(c))]
            if not take(" | ".join(cell.text.strip() for cell in cells)):
                break
        if truncated:
            break
    # A contract's total value sits in the footer often enough to matter.
    for section in document.sections if not truncated else []:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if not take(paragraph.text):
                    break
            if truncated:
                break
        if truncated:
            break

    return ExtractionResult(
        budget.text(),
        "truncated" if truncated else "complete",
        f"DOCX truncated to the first {MAX_EXTRACTED_CHARS:,} characters" if truncated else "",
    )


def _extract_xlsx(path: Path) -> ExtractionResult:
    try:
        import openpyxl  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional package
        return ExtractionResult("", "failed", f"XLSX extraction unavailable: {exc}")
    oversized = archive_too_large(path)
    if oversized is not None:
        return ExtractionResult("", "failed", archive_refusal_note("XLSX", oversized))
    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:
        return ExtractionResult("", "failed", f"XLSX could not be read (corrupt or not a real .xlsx): {exc}")
    budget = Budget()
    truncated = False
    try:
        for sheet in workbook.worksheets:
            if not budget.add(f"## Sheet: {sheet.title}"):
                truncated = True
                break
            for idx, row in enumerate(sheet.iter_rows(values_only=True)):
                if idx >= MAX_ROWS:
                    truncated = True
                    break
                if len(row) > MAX_COLUMNS:
                    truncated = True
                values = ["" if value is None else str(value) for value in row[:MAX_COLUMNS]]
                if not any(value.strip() for value in values):
                    continue
                if not budget.add(" | ".join(values)):
                    truncated = True
                    break
            if budget.exhausted:
                break
    finally:
        workbook.close()
    return ExtractionResult(
        budget.text(),
        "truncated" if truncated else "complete",
        limit_note("XLSX") if truncated else "",
    )


def _extract_pdf(path: Path) -> ExtractionResult:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional package
        return ExtractionResult("", "failed", f"PDF extraction unavailable: {exc}")
    budget = Budget()
    truncated = False
    try:
        reader = PdfReader(str(path))
        for page in reader.pages[:MAX_PDF_PAGES]:
            if not budget.add(page.extract_text() or ""):
                truncated = True
                break
        page_count = len(reader.pages)
    except Exception as exc:
        return ExtractionResult("", "failed", f"PDF could not be read (corrupt, encrypted, or scanned without OCR): {exc}")
    if page_count > MAX_PDF_PAGES:
        truncated = True
    return ExtractionResult(
        budget.text(),
        "truncated" if truncated else "complete",
        f"PDF limited to the first {MAX_PDF_PAGES} pages and {MAX_EXTRACTED_CHARS:,} characters" if truncated else "",
    )


def _clamped_result(text: str, kind: str) -> ExtractionResult:
    clamped, truncated = clamp(text)
    return ExtractionResult(
        clamped,
        "truncated" if truncated else "complete",
        f"{kind} truncated to the first {MAX_EXTRACTED_CHARS:,} characters" if truncated else "",
    )


def extract_text(path: Path) -> ExtractionResult:
    ext = path.suffix.lower()
    if ext in {".md", ".txt"}:
        return _clamped_result(read_text(path), "Text")
    if ext == ".csv":
        return _extract_csv(path)
    if ext == ".json":
        return _extract_json(path)
    if ext in {".html", ".htm"}:
        return _clamped_result(_strip_html(read_text(path)), "HTML")
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    return ExtractionResult("", "failed", "Unsupported file type")
